import os
import glob
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

def load_documents_from_folder(folder_path):
    """Loads text from .docx and .txt files in a folder."""
    documents = []
    if not os.path.exists(folder_path):
        print(f"Folder not found: {folder_path}")
        return documents

    files = glob.glob(os.path.join(folder_path, "**/*.*"), recursive=True)
    for path in files:
        if not os.path.isfile(path):
            continue
        
        content = ""
        try:
            if path.endswith(".docx"):
                import docx
                doc = docx.Document(path)
                content = "\n".join([p.text for p in doc.paragraphs])
            elif path.endswith(".pdf"):
                from pypdf import PdfReader
                reader = PdfReader(path)
                content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif path.endswith(".txt") or path.endswith(".md"): # Basic text support
                with open(path, "r", encoding="utf-8", errors="ignore") as file:
                    content = file.read()
            
            if content.strip():
                # Store relative path for cleaner metadata
                rel_path = os.path.relpath(path, start=os.getcwd())
                documents.append(Document(page_content=content, metadata={"source": rel_path}))
                print(f"Loaded: {rel_path}")
        except Exception as e:
            print(f"Error reading {path}: {e}")
    
    return documents

def main():
    # 1. Load environment variables
    load_dotenv()
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("Error: OPENAI_API_KEY not found.")
        return

    # 2. Load History Data
    history_folder = "history"
    print(f"Loading documents from '{history_folder}'...")
    docs = load_documents_from_folder(history_folder)
    
    retriever = None
    if docs:
        print(f"Vectorizer: Found {len(docs)} documents. Splitting and Indexing...")
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(docs)
        
        embeddings = OpenAIEmbeddings(api_key=api_key)
        vector_store = FAISS.from_documents(splits, embeddings)
        # k=3 for top 3 relevant chunks
        retriever = vector_store.as_retriever(search_kwargs={"k": 3})
        print("Vector Store initialized.")
    else:
        print("No documents found in history folder. Running in pure chat mode.")

    # 3. Initialize LLM & Chain
    llm = ChatOpenAI(model="gpt-4", api_key=api_key)
    
    # Prompt Template
    template = """You are a helpful assistant. Use the following context to answer the question.
If the answer is not in the context, say you don't know, but try to be helpful.

Context:
{context}

Question:
{question}
"""
    prompt = ChatPromptTemplate.from_template(template)
    
    def format_docs(docs):
        return "\n\n".join(f"[Source: {d.metadata['source']}]\n{d.page_content}" for d in docs)

    print("-" * 30)
    print("Ready! Type 'exit' to quit.")

    while True:
        user_input = input("\nYou: ")
        if user_input.lower() in ["exit", "quit"]:
            break
        if not user_input.strip():
            continue

        try:
            if retriever:
                # RAG Flow
                chain = (
                    {"context": retriever | format_docs, "question": RunnablePassthrough()}
                    | prompt
                    | llm
                    | StrOutputParser()
                )
                print("Thinking (searching history)...")
                response = chain.invoke(user_input)
            else:
                # Pure Chat Flow
                response = llm.invoke(user_input).content

            print(f"AI: {response}")
        except Exception as e:
            print(f"Error: {e}")

if __name__ == "__main__":
    main()
