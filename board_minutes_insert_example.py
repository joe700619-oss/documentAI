# 董事會議事錄 - 使用 insert 方式的示例 (逐列增加)
# 這個文件展示如何將 board_minutes_table 改為使用 insert 方式,逐列增加表格

# 1. Data 格式範例 - 每一列代表一個項目
board_minutes_data = {
    "tabel_ANCHOR": "[TABLE_BOARD_MINUTES_ANCHOR]",
    "column_widths": [1.5, 2.0, 14.5],  # 項目、案由/說明/決議、內容
    "num_header_rows": 0,  # 不需要表頭
    "agenda_rows": [
        {"項目": "1", "案由": "訂定公司章程案"},
        {"項目": "", "說明": "依公司法第129條規定,擬定公司章程,如附章程。"},
        {"項目": "", "決議": "經主席徵詢全體發起人無異議照案通過。"},
        {"項目": "2", "案由": "選任董事案"},
        {"項目": "", "說明": "選任董事3人。"},
        {"項目": "", "決議": "經主席徵詢全體董事無異議通過。"}
    ]
}

# 2. 需要添加到 test_script.py 的函式

"""
# 從 table_test.py 複製這些工具函式到 test_script.py:

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

def get_alignment(align_str):
    if not align_str:
        return WD_ALIGN_PARAGRAPH.LEFT
    return ALIGNMENT_MAP.get(align_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)

def cm_to_twips(cm):
    return int(cm * 567)

def set_col_widths(table, widths_cm):
    table.allow_autofit = False
    tbl = table._tbl
    tblGrid = tbl.tblGrid
    
    for i, cm_val in enumerate(widths_cm):
        twips = cm_to_twips(cm_val)
        if i < len(tblGrid.gridCol_lst):
            gridCol = tblGrid.gridCol_lst[i]
            gridCol.w = twips
        for row in table.rows:
            if i < len(row.cells):
                row[i].width = Cm(cm_val)

def set_cell_format(cell, text, size_pt, is_bold=False, alignment=None):
    p = cell.paragraphs[0]
    p.clear() 
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = is_bold
    if alignment:
        p.alignment = alignment

def insert_table_at_anchor(document, anchor_text, num_cols, num_rows=1):
    # num_rows 改為初始行數,因為不需要表頭,所以從1開始
    placeholder_paragraph = None
    for paragraph in document.paragraphs:
        if anchor_text in paragraph.text:
            placeholder_paragraph = paragraph
            break
            
    if not placeholder_paragraph:
        print(f"警告：未找到錨點 '{anchor_text}'")
        return None
    
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    t_element = table._element
    t_element.getparent().remove(t_element)
    
    p_element = placeholder_paragraph._element
    p_parent = p_element.getparent()
    p_element.addprevious(t_element)
    p_parent.remove(p_element)
    
    return table

# 3. 新增議程表格填充函式 - 逐列增加
def populate_board_minutes_table(table, table_data, num_cols):
    column_widths = table_data.get("column_widths")
    agenda_rows = table_data.get("agenda_rows", [])
    
    # 設置第一行(已經存在)
    if len(agenda_rows) > 0:
        first_row = agenda_rows[0]
        cells = table.rows[0].cells
        
        # 設置欄寬
        for i, w in enumerate(column_widths):
            if i < len(cells):
                cells[i].width = Cm(w)
        
        # 填充第一行內容
        set_cell_format(cells[0], first_row.get("項目", ""), 12, False, get_alignment("center"))
        
        # 判斷是案由、說明還是決議
        if "案由" in first_row:
            set_cell_format(cells[1], "案由", 12, True, get_alignment("left"))
            set_cell_format(cells[2], first_row.get("案由", ""), 12, False, get_alignment("left"))
        elif "說明" in first_row:
            set_cell_format(cells[1], "說明", 12, True, get_alignment("left"))
            set_cell_format(cells[2], first_row.get("說明", ""), 12, False, get_alignment("left"))
        elif "決議" in first_row:
            set_cell_format(cells[1], "決議", 12, True, get_alignment("left"))
            set_cell_format(cells[2], first_row.get("決議", ""), 12, False, get_alignment("left"))
    
    # 設置欄寬
    set_col_widths(table, column_widths)
    
    # 逐列增加剩餘的行
    for i in range(1, len(agenda_rows)):
        row_data = agenda_rows[i]
        
        # 新增一行
        new_row = table.add_row().cells
        
        # 設置欄寬
        for j, w in enumerate(column_widths):
            if j < len(new_row):
                new_row[j].width = Cm(w)
        
        # 填充內容
        set_cell_format(new_row[0], row_data.get("項目", ""), 12, False, get_alignment("center"))
        
        # 判斷是案由、說明還是決議
        if "案由" in row_data:
            set_cell_format(new_row[1], "案由", 12, True, get_alignment("left"))
            set_cell_format(new_row[2], row_data.get("案由", ""), 12, False, get_alignment("left"))
        elif "說明" in row_data:
            set_cell_format(new_row[1], "說明", 12, True, get_alignment("left"))
            set_cell_format(new_row[2], row_data.get("說明", ""), 12, False, get_alignment("left"))
        elif "決議" in row_data:
            set_cell_format(new_row[1], "決議", 12, True, get_alignment("left"))
            set_cell_format(new_row[2], row_data.get("決議", ""), 12, False, get_alignment("left"))

# 4. 在主程式中使用
def process_board_minutes(doc, table_data):
    anchor_text = table_data.get("tabel_ANCHOR")
    num_cols = len(table_data.get("column_widths"))
    
    # 創建初始表格(1行,因為不需要表頭)
    inserted_table = insert_table_at_anchor(doc, anchor_text, num_cols, num_rows=1)
    if inserted_table:
        populate_board_minutes_table(inserted_table, table_data, num_cols)
        print(f"✅ 董事會議事錄表格已插入並填充完成")

# 5. 在 test_script.py 的 generate_document() 中使用:
'''
# 在處理文件時,檢查是否為董事會議事錄
if "董事會議事錄" in doc_name:
    # 使用新的 insert 方式
    board_minutes_data = data["table_data"].get("board_minutes_table")
    if board_minutes_data:
        process_board_minutes(doc, board_minutes_data)
else:
    # 其他文件使用原有的處理方式
    replace_all_placeholders(doc, replacements)
'''
"""

# 完整的 data 格式範例
complete_data_example = {
    "basic_info": {
        # ... 保持原有的 basic_info
    },
    "table_data": {
        "board_minutes_table": {
            "tabel_ANCHOR": "[TABLE_BOARD_MINUTES_ANCHOR]",
            "column_widths": [1.5, 2.0, 14.5],
            "num_header_rows": 0,
            "agenda_rows": [
                {"項目": "1", "案由": "訂定公司章程案"},
                {"項目": "", "說明": "依公司法第129條規定,擬定公司章程,如附章程。"},
                {"項目": "", "決議": "經主席徵詢全體發起人無異議照案通過。"},
                {"項目": "2", "案由": "選任董事案"},
                {"項目": "", "說明": "選任董事3人。"},
                {"項目": "", "決議": "經主席徵詢全體董事無異議通過。"}
            ]
        },
        # ... 其他表格保持不變
    }
}

print("✅ 示例文件已創建")
print("📋 Data 格式:")
print("   - 不需要表頭 (num_header_rows: 0)")
print("   - 使用 agenda_rows 陣列,每個元素代表一列")
print("   - 每列包含: 項目、案由/說明/決議")
print("\n📝 使用方式:")
print("   1. 複製上述函式到 test_script.py")
print("   2. 修改 data['table_data']['board_minutes_table'] 為新格式")
print("   3. 在處理董事會議事錄時調用 process_board_minutes()")
