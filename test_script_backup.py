import json
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL # 確保匯入，用於 set_cell_format 中的垂直置中

# --- 全域對齊方式對應表 ---
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

def get_alignment(align_str):
    """將字串轉換為 WD_ALIGN_PARAGRAPH 常數，預設為左對齊"""
    if not align_str:
        return WD_ALIGN_PARAGRAPH.LEFT
    return ALIGNMENT_MAP.get(align_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)

# --- 工具函數 1：cm 轉 twips ---
def cm_to_twips(cm):
    """將公分轉換為 Word 的內部單位 twips (二十份點)"""
    return int(cm * 567)

# --- 工具函數 2：設定整張表格欄寬（真正生效） ---
def set_col_widths(table, widths_cm):
    """
    通過操作底層 XML，強制設定表格所有欄位的寬度。
    """
    # 禁止自動調整欄寬
    table.allow_autofit = False
    
    tbl = table._tbl
    tblGrid = tbl.tblGrid

    if len(tblGrid.gridCol_lst) < len(widths_cm):
        # 由於我們在 insert_table_at_anchor 中已經創建了正確的欄位數，
        # 通常這裡不會出錯。
        pass

    for i, cm_val in enumerate(widths_cm):
        twips = cm_to_twips(cm_val)
        # 設置 gridCol 元素的寬度 (如果存在)
        if i < len(tblGrid.gridCol_lst):
            gridCol = tblGrid.gridCol_lst[i]
            gridCol.w = twips
        
        # 保持額外設定：設定現有列的儲存格寬度
        for row in table.rows:
            if i < len(row.cells):
                row.cells[i].width = Cm(cm_val)

# --- 工具函數 4：批量設定單元格格式 ---
def set_cell_format(cell, text, size_pt, is_bold=False, alignment=None):
    """
    設定單元格的文字、字體大小、粗體和對齊方式。
    確保格式正確應用，使用 p.clear() 和 p.add_run()。
    """
    p = cell.paragraphs[0]
    p.clear() 
    
    # 設置文字 Run 格式
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = is_bold
    
    # 設置對齊方式 (如果指定了)
    if alignment:
        p.alignment = alignment

# --- 工具函數 3：錨點插入（修正為接受動態行數） ---
def insert_table_at_anchor(document, anchor_text, num_cols, num_header_rows):
    placeholder_paragraph = None
    
    # 1. 尋找錨點段落
    for paragraph in document.paragraphs:
        if anchor_text in paragraph.text:
            placeholder_paragraph = paragraph
            break
            
    if not placeholder_paragraph:
        print(f"警告：未找到錨點 '{anchor_text}'，表格將插入在文件末尾。")
        return document.add_table(rows=num_header_rows, cols=num_cols) 
    
    # 2. 創建表格 (rows=num_header_rows, cols=num_cols)
    table = document.add_table(rows=num_header_rows, cols=num_cols)
    table.style = 'register_table'
    
    # 3. 獲取表格的 XML 元素
    t_element = table._element
    
    # 4. 刪除表格在文件末尾的預設位置 (關鍵步驟)
    t_element.getparent().remove(t_element)
    
    # 5. 將表格 XML 元素插入到錨點 XML 元素的前面
    p_element = placeholder_paragraph._element
    p_parent = p_element.getparent()
    p_element.addprevious(t_element)

    # 6. 刪除錨點段落
    p_parent.remove(p_element)
    
    return table


# ----------------------------------------------------
# --- 通用表頭構建函式 (核心抽象) ---
# ----------------------------------------------------

# --- 工具函數 5a：構建 2 行表頭 (為 "所營事業" 設計) ---
def build_2_row_header(table, table_data, num_cols):
    """根據 header_structure 數據，構建 2 行表頭。"""
    header_structure = table_data["header_structure"]

    hdr1 = table.rows[0].cells
    hdr2 = table.rows[1].cells
    
    # Hdr1: 總標題
    hdr1[0].merge(hdr1[num_cols - 1])
    title_data = header_structure["title"]
    fmt = title_data["format"]
    set_cell_format(hdr1[0], title_data["text"], fmt["size_pt"], fmt["is_bold"], get_alignment(fmt["alignment"]))

    # Hdr2: 設置主欄位名稱 (迴圈處理所有欄位)
    row_2_fields = header_structure["row_2_fields"]
    for i, cell_data in enumerate(row_2_fields):
        fmt = cell_data["format"]
        set_cell_format(
            hdr2[i], 
            cell_data["text"], 
            fmt["size_pt"], 
            fmt["is_bold"], 
            get_alignment(fmt["alignment"])
        )
    
    set_col_widths(table, table_data.get("column_widths"))

# --- 補回：工具函數 5b：構建三層複雜表頭 (供 build_table_header 呼叫) ---
def build_3_row_header(table, table_data, num_cols):
    """
    通用函數：根據 header_structure 數據，構建帶有垂直/水平合併的三層表頭。
    """
    header_structure = table_data["header_structure"]

    hdr1 = table.rows[0].cells
    hdr2 = table.rows[1].cells
    hdr3 = table.rows[2].cells
    
    # 1. Hdr1: 總標題 (合併所有欄位)
    hdr1[0].merge(hdr1[num_cols - 1])
    title_data = header_structure["title"]
    fmt = title_data["format"]
    set_cell_format(
        hdr1[0], 
        title_data["text"], 
        fmt["size_pt"], 
        fmt["is_bold"], 
        get_alignment(fmt["alignment"])
    )

    # 2. Hdr2 & Hdr3: 垂直合併「編號」欄
    hdr2[0].merge(hdr3[0])
    id_data = header_structure["id_column"]
    fmt_id = id_data["format"]
    set_cell_format(
        hdr2[0], 
        id_data["text"], 
        fmt_id["size_pt"], 
        fmt_id["is_bold"], 
        get_alignment(fmt_id["alignment"])
    )

    # 3. Hdr2: 設置主欄位名稱 (迴圈處理其餘欄位)
    for i, cell_data in enumerate(header_structure["row_2_fields"]):
        cell_index = i + 1  # 從第二個單元格開始 (索引 1)
        fmt = cell_data["format"]
        set_cell_format(
            hdr2[cell_index], 
            cell_data["text"], 
            fmt["size_pt"], 
            fmt["is_bold"], 
            get_alignment(fmt["alignment"])
        )
        
    # 4. Hdr3: 水平合併「地址」欄
    addr_data = header_structure["row_3_address"]
    if num_cols > 1:
        hdr3[1].merge(hdr3[num_cols - 1])
    
    fmt_addr = addr_data["format"]
    set_cell_format(
        hdr3[1], 
        addr_data["text"], 
        fmt_addr["size_pt"], 
        fmt_addr["is_bold"], 
        get_alignment(fmt_addr["alignment"])
    )
    
    set_col_widths(table, table_data.get("column_widths"))


# --- 工具函數 5c：通用表頭構建入口 ---
def build_table_header(table, table_data, num_cols, num_header_rows):
    """根據行數呼叫對應的表頭構建邏輯。"""
    if num_header_rows == 3:
        build_3_row_header(table, table_data, num_cols) 
    elif num_header_rows == 2:
        build_2_row_header(table, table_data, num_cols)
    else:
        print(f"警告: 不支援 {num_header_rows} 行表頭的結構。")


# --- 填充函式 6：萬用表格填充 (單一入口) ---
def populate_universal_table(table, table_data, num_cols, num_header_rows):
    """
    通用填充函式：處理表頭、數據列、欄寬設定和地址行合併。
    """
    column_widths = table_data.get("column_widths")
    
    # 1. 設置表頭
    build_table_header(table, table_data, num_cols, num_header_rows)
    
    # 2. 填充數據行
    data_list = table_data.get(table_data["data_source_key"], [])
    data_mapping = table_data["data_row_mapping"]
    addr_config = table_data.get("address_config")

    for item in data_list:
        
        # 2a. 第一列：基本資料
        row = table.add_row().cells
        
        # 確保新的數據行應用寬度
        for i, w in enumerate(column_widths):
            if i < len(row):
                row[i].width = Cm(w)

        # 根據 data_row_mapping 填充所有欄位 (透過迴圈實現簡潔)
        for mapping in data_mapping:
            col_index = mapping["col_index"]
            source_key = mapping["source_key"]
            fmt = mapping["format"]
            default_val = mapping.get("default_val", "")
            
            text = item.get(source_key, default_val)
            
            set_cell_format(
                row[col_index], 
                text, 
                fmt["size_pt"], 
                fmt["is_bold"], 
                get_alignment(fmt["alignment"])
            )
            
        # 2b. 第二列：地址/所在地 (如果存在配置且數據中有地址)
        if addr_config and item.get(addr_config["source_key"]):
            address = item[addr_config["source_key"]]
            
            addr_row = table.add_row().cells
            for i, w in enumerate(column_widths):
                if i < len(addr_row):
                    addr_row[i].width = Cm(w)
                    
            # 垂直合併 ID 欄 (固定為第 0 欄)
            row[0].merge(addr_row[0]) 
            
            # 設置地址單元格和水平合併
            start_col = addr_config["start_col"]
            addr_cell = addr_row[start_col]
            
            # 水平合併從 start_col 到最後一欄
            addr_cell.merge(addr_row[num_cols - 1]) 

            fmt = addr_config["format"]
            set_cell_format(
                addr_cell, 
                address, 
                fmt["size_pt"], 
                fmt["is_bold"], 
                get_alignment(fmt["alignment"])
            )

# ----------------------------------------------------
# --- 移除冗餘的 TABLE_POPULATORS (因為不再需要) ---
# ----------------------------------------------------
# TABLE_POPULATORS = { ... } # 已移除
# ----------------------------------------------------


# ==========================================================
# 最終執行區塊
# ==========================================================

# Embedded data (已移除 director_table 中的多餘逗號)
data = {
"director_table": {
        "document_title": "董事、監察人名單",
        "tabel_ANCHOR": "[TABLE_DIRECTORS_ANCHOR]",
        "column_widths": [0.75, 2.01, 5.61, 6.03, 3.49],
        "num_header_rows": 3,
        "data_source_key": "directors", 
        "data_row_mapping": [
            {"col_index": 0, "source_key": "id", "format": {"size_pt": 12, "is_bold": False, "alignment": "center"}},
            {"col_index": 1, "source_key": "position", "default_val": "", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 2, "source_key": "name", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 3, "source_key": "id_number", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 4, "source_key": "shares", "format": {"size_pt": 12, "is_bold": False, "alignment": "right"}},
        ],
        "address_config": {
            "source_key": "address",
            "format": {"size_pt": 12, "is_bold": False, "alignment": "left"},
            "start_col": 1,
        },
        "header_structure": {
            "title": {
                "text": "董事、監察人名單",
                "format": {"size_pt": 16, "is_bold": True, "alignment": "center"}
            },
            "id_column": {
                "text": "編號",
                "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}
            },
            "row_2_fields": [
                {"text": "職稱", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "姓名(或法人名稱)", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "身分證號(或法人統一編號)", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "持有股份(股)", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}}
            ],
            "row_3_address": {
                "text": "(郵遞區號)住所或居所(或法人所在地)",
                "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}
            }
        },
        "directors": [
            {
                "id": "001", "position": "董事長", "name": "李耿佑", "id_number": "F128873285", "shares": "20,000",
                "address": "(220) 新北市板橋區湳興里 4 鄰南雅西路二段 7 巷 18 之 4 號"
            },
            {
                "id": "002", "position": "監察人", "name": "李耿甫", "id_number": "A123456789", "shares": "0",
                "address": "台北市信義區忠孝東路一段1號"
            }
        ]
    },
    "manager_table": {
        "document_title": "經理人名單",
        "tabel_ANCHOR": "[TABLE_MANAGER_ANCHOR]",
        "column_widths": [0.75, 5.42, 5.37, 6.38],
        "num_header_rows": 3,
        "data_source_key": "managers",
        "data_row_mapping": [
            {"col_index": 0, "source_key": "id", "format": {"size_pt": 12, "is_bold": False, "alignment": "center"}},
            {"col_index": 1, "source_key": "name", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 2, "source_key": "id_number", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 3, "source_key": "onboarding_date", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
        ],
        "address_config": {
            "source_key": "address",
            "format": {"size_pt": 12, "is_bold": False, "alignment": "left"},
            "start_col": 1,
        },
        "header_structure": {
            "title": {
                "text": "經理人名單",
                "format": {"size_pt": 16, "is_bold": True, "alignment": "center"}
            },
            "id_column": {
                "text": "編號",
                "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}
            },
            "row_2_fields": [
                {"text": "姓名", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "身分證號", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "到職日期(年月日)", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
            ],
            "row_3_address": {
                "text": "(郵遞區號)住所或居所",
                "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}
            }
        },
        "managers": [
            {
                "id": "001", "name": "joe", "id_number": "F128873285", "onboarding_date": "114年11月29日",
                "address": "(220) 新北市板橋區湳興里 4 鄰南雅西路二段 7 巷 18 之 4 號"
            }
        ]
    },
    "representative_of_the_company_table": {
        "document_title": "所代表法人",
        "tabel_ANCHOR": "[TABLE_REPRESENTATIVE_OF_THE_COMPANY_ANCHOR]",
        "column_widths": [0.75, 3.47, 7, 6.68],
        "num_header_rows": 3,
        "data_source_key": "representative_of_the_company",
        "data_row_mapping": [
            {"col_index": 0, "source_key": "id", "format": {"size_pt": 12, "is_bold": False, "alignment": "center"}},
            {"col_index": 1, "source_key": "serialNumber", "default_val": "", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 2, "source_key": "name", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 3, "source_key": "id_number", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
        ],
        "address_config": {
            "source_key": "address",
            "format": {"size_pt": 12, "is_bold": False, "alignment": "left"},
            "start_col": 1,
        },
        "header_structure": {
            "title": {
                "text": "所代表法人",
                "format": {"size_pt": 16, "is_bold": True, "alignment": "center"}
            },
            "id_column": {
                "text": "序號",
                "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}
            },
            "row_2_fields": [
                {"text": "董監事編號", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "所代表法人名稱", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "法人統一編號", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
            ],
            "row_3_address": {
                "text": "(郵遞區號)法人所在地",
                "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}
            }
        },
        "representative_of_the_company": [
            {
                "id": "001", "serialNumber": "01~01", "name": "joe", "id_number": "F128873285",
                "address": "(220) 新北市板橋區湳興里 4 鄰南雅西路二段 7 巷 18 之 4 號"
            }
        ]
    },
    "business_scope_table": {
        "document_title": "所營事業",
        "tabel_ANCHOR": "[TABLE_BUSINESS_SCOPE_ANCHOR]",
        "column_widths": [1.21, 3.7, 13.18],
        "num_header_rows": 2,
        "data_source_key": "business_scope",
        "data_row_mapping": [
            {"col_index": 0, "source_key": "id", "format": {"size_pt": 12, "is_bold": False, "alignment": "center"}},
            {"col_index": 1, "source_key": "business_code", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
            {"col_index": 2, "source_key": "business_scription", "format": {"size_pt": 12, "is_bold": False, "alignment": "left"}},
        ],
        "address_config": None,
        "header_structure": {
            "title": {
                "text": "所營事業",
                "format": {"size_pt": 16, "is_bold": True, "alignment": "center"}
            },
            "row_2_fields": [
                {"text": "編號", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "代碼", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
                {"text": "營業項目說明", "format": {"size_pt": 12, "is_bold": True, "alignment": "center"}},
            ],
        },
        "business_scope": [
            {
                "id": "001", "business_code": "J101010", "business_scription": "建築物清潔服務業", 
            },
            {
                "id": "002", "business_code": "JA03010", "business_scription": "洗衣業", 
            },
            {
                "id": "003", "business_code": "JZ99990", "business_scription": "未分類其他服務業", 
            },
        ]
    }
}


# 載入文件
document = Document(r"C:\Users\joe70\PythonProject\documentAI\templates\設立登記表.docx")

# --------------------------------------
# 核心執行：迭代所有表格數據並插入/填充
# --------------------------------------

for table_key, table_data in data.items():
    
    anchor_text = table_data.get("tabel_ANCHOR")
    num_cols = len(table_data.get("column_widths"))
    num_header_rows = table_data.get("num_header_rows", 3)
    
    # 1. 插入表格到指定位置 (傳遞正確的欄位數和表頭行數)
    inserted_table = insert_table_at_anchor(document, anchor_text, num_cols, num_header_rows)

    # 2. 統一填充表格
    if inserted_table:
        # 呼叫萬用填充函式，由其處理表頭構建和數據行填充
        populate_universal_table(inserted_table, table_data, num_cols, num_header_rows)
        
# 輸出結果
output_dir = "output"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "output_v2.docx")
document.save(output_path)
print(f"文件已產生：{output_path}")