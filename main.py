
import os
import json
import re
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import SystemMessage
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from getbasicInformationfromMOEA import BasicInformationAPI

load_dotenv()

# Since the user asked to send a message to the LLM model, 
# I will implement a class that simulates or interfaces with an LLM.
# If an API key is present, it can use it. Otherwise, it will mock or allow the user to see the prompt.
# For this task, I will use a Mock or a placeholder unless an API library is requested.
# However, the previous main.py had imports for langchain_openai. 
# I will attempt to reuse or structure it such that it CAN use it if available, 
# but specifically addressing the NEW requirements.

# NOTE: The user's previous main.py was a mess of merge conflicts (HEAD vs commit hash).
# I will completely rewrite it to follow the NEW requirements strictly.

class DocumentAIWorkflow:
    def __init__(self):
        self.api = BasicInformationAPI()
        self.data_schema_path = "data.json"
        self.test_data_path = "test_data.json"
        self.history_dir = "history"
        self.history_cases_dir = "history_cases"
        self.vector_store_path = "faiss_index"
        self.api_key = os.environ.get("OPENAI_API_KEY")

    def load_json(self, path):
        if not os.path.exists(path):
            return {}
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

    def load_documents_from_folder(self, folder_path):
        """Loads text from files in a folder and returns a list of Document objects."""
        documents = []
        if not os.path.exists(folder_path):
            return documents

        for f in os.listdir(folder_path):
            path = os.path.join(folder_path, f)
            if os.path.isfile(path):
                content = ""
                try:
                    if path.endswith(".docx"):
                        import docx
                        doc = docx.Document(path)
                        content = "\n".join([p.text for p in doc.paragraphs])
                    elif path.endswith(".pdf"):
                        from pypdf import PdfReader
                        reader = PdfReader(path)
                        content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                    else:
                        with open(path, "r", encoding="utf-8", errors="ignore") as file:
                            content = file.read()
                    
                    if content.strip():
                        documents.append(Document(page_content=content, metadata={"source": f}))
                except Exception as e:
                    print(f"Error reading {path}: {e}")
        return documents

    def setup_vector_store(self):
        """
        Loads documents, splits them, and creates/loads a vector store.
        """
        print("Loading history documents...")
        history_docs = self.load_documents_from_folder(self.history_dir)
        cases_docs = self.load_documents_from_folder(self.history_cases_dir)
        all_docs = history_docs + cases_docs

        if not all_docs:
            print("No history documents found to vectorize.")
            return None

        # text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(all_docs)

        print("Creating embeddings and vector store...")
        if self.api_key:
            embeddings = OpenAIEmbeddings(api_key=self.api_key)
            vector_store = FAISS.from_documents(splits, embeddings)
            return vector_store
        else:
            print("No API Key available for embeddings.")
            return None

    def run_rag_chain(self, user_input, authoritative_data, schema):
        vector_store = self.setup_vector_store()
        # Reduce k to avoid blowing up context window if chunks are dense
        retriever = vector_store.as_retriever(search_kwargs={"k": 2}) if vector_store else None

        facts_str = json.dumps(authoritative_data.get("company_facts_authoritative", {}), ensure_ascii=False, indent=2)
        policy_str = json.dumps(authoritative_data.get("field_policy", {}), ensure_ascii=False, indent=2)
        schema_str = json.dumps(schema, ensure_ascii=False, indent=2)

        system_template = """You are an intelligent data processing assistant.
Your goal is to output a JSON object that matches the EXACT structure of the provided 'Target Schema'.

### Sources
1. **Authoritative Data** (High Priority, heavily restricted):
{facts}

2. **Field Policy**:
{policy}
   - If a field is "locked", you MUST use the value from 'Authoritative Data' *UNLESS* the 'User Request' explicitly asks to change it.
   - If a field is "editable", you should fill it based on 'History Context' or 'User Request'.
   - If a field is NOT in 'Authoritative Data' but is in 'Target Schema', treat it as "editable".

3. **History Context**:
{context}

4. **Target Schema**:
{schema}

### Instructions
- Generate a valid JSON object matching 'Target Schema'.
- Fill "locked" fields from 'Authoritative Data'.
- modify "locked" fields ONLY if the user request specifically contradicts them.
- Fill "editable" fields intelligently using the User Request and History Context.
- Return ONLY the JSON string, no markdown formatting.
"""
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_template),
            ("human", "{question}")
        ])

        if self.api_key:
            llm = ChatOpenAI(model="gpt-4", api_key=self.api_key)
            
            def format_docs(docs):
                return "\n\n".join(doc.page_content for doc in docs)

            # RAG Chain
            if retriever:
                rag_chain = (
                    {"context": retriever | format_docs, "question": RunnablePassthrough(), 
                     "facts": lambda x: facts_str, "policy": lambda x: policy_str, "schema": lambda x: schema_str}
                    | prompt
                    | llm
                    | StrOutputParser()
                )
            else:
                # Fallback without retriever
                rag_chain = (
                    {"context": lambda x: "No history context available.", "question": RunnablePassthrough(),
                     "facts": lambda x: facts_str, "policy": lambda x: policy_str, "schema": lambda x: schema_str}
                    | prompt
                    | llm
                    | StrOutputParser()
                )

            print("\n[Simulated] Sending prompt to LLM (via Chain)...")
            return rag_chain.invoke(user_input)
        else:
             return json.dumps({"error": "No API Key"}, ensure_ascii=False)

    def run(self):
        # 1. Get User Input
        user_input = input("Please enter your request (e.g. 'Change chairman to Zhang San'): ")
        
        # 2. Run getbasicInformationfromMOEA logic
        print("Fetching authoritative data...")
        try:
            from getbasicInformationfromMOEA import main as fetch_main
            fetch_main()
        except Exception as e:
            print(f"Warning: Could not run fetch script directly: {e}")
        
        authoritative_data = self.load_json(self.test_data_path)
        schema = self.load_json(self.data_schema_path)

        # 3. Process with RAG
        result_json_str = self.run_rag_chain(user_input, authoritative_data, schema)
        
        # 4. Save Final Output
        try:
            cleaned_str = result_json_str.replace("```json", "").replace("```", "").strip()
            result_data = json.loads(cleaned_str)
            
            output_file = "final_llm_output.json"
            with open(output_file, "w", encoding="utf-8") as f:
                json.dump(result_data, f, ensure_ascii=False, indent=4)
            print(f"\nFinal data saved to {output_file}")
            
        except json.JSONDecodeError:
            print("LLM did not return valid JSON.")
            print("Raw output:", result_json_str)

if __name__ == "__main__":
    app = DocumentAIWorkflow()
    app.run()
