# 測試關鍵字替換儲存格內容 - 從 data.json 讀取設定
from docx import Document
from pathlib import Path
import json


def format_replacement_content(table_config):
    """
    根據 table_config 的結構，格式化替換內容
    """
    target_text = table_config.get("target_start_text", "")
    
    # 判斷資料類型並格式化
    if "business_items" in table_config:
        items = table_config["business_items"]
        # 檢查是營業項目還是章程修訂日期
        if items and "code" in items[0]:
            # 營業項目格式: id. code name
            lines = [target_text]
            for item in items:
                lines.append(f"{item['id']}. {item['code']} {item['name']}")
            return "\n".join(lines)
        elif items and "edition_date" in items[0]:
            # 章程修訂日期格式
            lines = [target_text]
            for item in items:
                lines.append(f"{item['id']}. {item['edition_date']}")
            return "\n".join(lines)
    
    if "shareholders" in table_config:
        items = table_config["shareholders"]
        lines = [target_text]
        for item in items:
            lines.append(f"{item['name']}：{item['capital']}")
        return "\n".join(lines)
    
    return target_text


def find_and_replace_by_config(document, table_config):
    """
    根據 table_config 中的 target_start_text 找到儲存格並替換內容
    """
    target_start_text = table_config.get("target_start_text", "")
    if not target_start_text:
        return False, "No target_start_text defined"
    
    # 格式化替換內容
    new_content = format_replacement_content(table_config)
    
    # 遍歷所有表格
    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 移除冒號差異，做更寬鬆的比對
                target_clean = target_start_text.rstrip("：:")
                if cell_text.startswith(target_clean):
                    print(f"  [FOUND] Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}")
                    print(f"    Original: {cell_text[:50]}...")
                    
                    # 清除原有內容
                    for paragraph in cell.paragraphs:
                        paragraph.clear()
                    
                    # 寫入新內容
                    cell.text = new_content
                    print(f"    Replaced!")
                    return True, "Replacement successful"
    
    return False, f"Text starting with '{target_start_text}' not found"


def main():
    # 設定路徑
    base_dir = Path(r"c:\Users\joe70\PythonProject\documentAI")
    template_path = base_dir / "templates" / "公司章程.docx"
    output_path = base_dir / "output" / "公司章程.docx"
    data_path = base_dir / "data.json"
    
    print("=" * 60)
    print("Test: Find and Replace from data.json")
    print("=" * 60)
    
    # 讀取 data.json
    print(f"\nLoading: {data_path}")
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 取得所有含有 target_start_text 的表格設定
    table_data = data.get("tableData", {})
    target_tables = []
    for table_name, table_config in table_data.items():
        if "target_start_text" in table_config:
            target_tables.append({
                "name": table_name,
                "config": table_config
            })
    
    print(f"\nFound {len(target_tables)} tables with target_start_text:")
    for t in target_tables:
        print(f"  - {t['name']}: '{t['config']['target_start_text']}'")
    
    # 確認範本存在
    if not template_path.exists():
        print(f"\nERROR: Template not found: {template_path}")
        return
    
    # 載入文件
    print(f"\nLoading template: {template_path}")
    document = Document(template_path)
    
    # 執行每個替換
    print("\n" + "-" * 60)
    print("Processing replacements:")
    print("-" * 60)
    
    results = []
    for t in target_tables:
        print(f"\n[{t['name']}]")
        success, message = find_and_replace_by_config(document, t['config'])
        results.append({
            "name": t['name'],
            "success": success,
            "message": message
        })
    
    # 儲存文件
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    
    # 顯示結果
    print("\n" + "=" * 60)
    print("Results:")
    print("=" * 60)
    for r in results:
        status = "OK" if r['success'] else "SKIP"
        print(f"  [{status}] {r['name']}: {r['message']}")
    
    print(f"\nOutput saved to: {output_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()