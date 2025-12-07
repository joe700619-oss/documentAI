"""
檢查發起人名冊範本 - 完整變數列表
"""
from docx import Document
from pathlib import Path
import re
import json

template_path = Path(r"c:\Users\joe70\PythonProject\documentAI\templates\發起人名冊.docx")
data_path = Path(r"c:\Users\joe70\PythonProject\documentAI\data.json")
output_path = Path(r"c:\Users\joe70\PythonProject\documentAI\check_result.txt")

# Load data.json
with open(data_path, 'r', encoding='utf-8') as f:
    data = json.load(f)

# Build available context
context = {}
if 'basicInformation' in data:
    context.update(data['basicInformation'])
if 'tableData' in data:
    for table_name, table_content in data['tableData'].items():
        if isinstance(table_content, dict):
            for key, value in table_content.items():
                if key != 'document_title':
                    context[key] = value

lines = []
lines.append("=== Available in context ===")
for k in sorted(context.keys()):
    v = context[k]
    if isinstance(v, list):
        lines.append(f"  {k}: list ({len(v)} items)")
    else:
        lines.append(f"  {k}: {v}")

# Find template variables
doc = Document(template_path)
template_vars = set()

for para in doc.paragraphs:
    matches = re.findall(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*(?:\.[a-zA-Z_][a-zA-Z0-9_]*)?)\s*\}\}', para.text)
    template_vars.update(matches)
    
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            matches = re.findall(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*(?:\.[a-zA-Z_][a-zA-Z0-9_]*)?)\s*\}\}', cell.text)
            template_vars.update(matches)

lines.append("\n=== Template variables ===")
for v in sorted(template_vars):
    lines.append(f"  {v}")

# Check which are missing
lines.append("\n=== MISSING (will cause error) ===")
missing = []
for v in sorted(template_vars):
    if '.' in v:
        continue  # Skip object.attr (handled by loop)
    if v not in context:
        missing.append(v)
        lines.append(f"  {v}")

if not missing:
    lines.append("  (none)")

# Write result
with open(output_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print(f"Result written to {output_path}")
