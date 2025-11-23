import os 
import json 
import re 
import ast 
from docx import Document 
import os 
import json 
import re 
import ast 
from docx import Document 
from docx.shared import Pt 
from docx.oxml.ns import qn 
from docx.table import _Row # 引入內部類別，用於 Row 的 XML 複製
from copy import deepcopy # <-- 新增此行
# ... (其他 import 保持不變)

import registration_doc

# -----------------------------------------------------
# 設定路徑 (根據您的環境變數設置，請自行調整)
# -----------------------------------------------------
BASE_DIR = r"C:\Users\joe70\PythonProject\documentAI" 
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
TEMPLATE_FILE = os.path.join(TEMPLATE_DIR, "board_minutes_template.docx")


# -----------------------------------------------------
# 輔助函式 (A): 表格 Row 區塊複製
# -----------------------------------------------------
def _copy_row_block(source_table, start_index, count):
    """
    複製從 start_index 開始的 count 個 Row，並將它們追加到表格末尾。
    返回新追加的 Row 列表。
    """
    new_rows = []
    tbl = source_table._tbl
    
    # 獲取要複製的 Row 的 XML 元素
    rows_to_copy = [source_table.rows[start_index + j]._tr for j in range(count)]
    
    for tr in rows_to_copy:
        # 克隆 XML 元素 (深度複製以避免引用問題)
        new_tr = deepcopy(tr)
        # 將克隆的 XML 元素追加到表格的 XML 中
        tbl.append(new_tr)
        # 創建新的 docx Row 物件來管理這個 XML 元素
        new_row = _Row(new_tr, source_table)
        new_rows.append(new_row)
        
    return new_rows

# -----------------------------------------------------
# 輔助函式 (B): 根據表頭尋找表格
# -----------------------------------------------------
def find_table_by_header(document, header_text):
    """
    在文件中尋找第一個儲存格內容匹配 header_text 的表格。
    """
    for table in document.tables:
        if len(table.rows) > 0:
            # 檢查表格的第一個儲存格內容
            first_cell_text = table.rows[0].cells[0].text.strip()
            if first_cell_text == header_text:
                return table
    return None

# -----------------------------------------------------
# AI 回傳的 JSON (已包含 'agend' 列表)
# -----------------------------------------------------

#{{companyName}}, {{applyReason}},{{reservationNumber}},{{governmentfees}},{{companyId}},{{chairperson}}
#{{zipcode}},{{address}},{{change_type}},{{new_directors_table}},{{removal_directors_table}},{{directors_table}}



data = {

  "basic_info": {
    "companyName": "茉莉家事服務股份有限公司",
    "companyName_eng": "Jasmine Home Service Co., Ltd.",
    "chairperson": "李耿佑",
    "company_address": "臺北市中山區民權西路20號二樓",
    "zipcode": "104",
    "addressCounty": "臺北市",
    "parvalue": 10,
    "capitalAmount": "1,000,000",
    "totalshares": "100,000",
    "totalcapital": "1,000,000",
    "date": "113年 04 月 22 日",
    "record_keeper": "李耿甫",
    "meeting_place": "本公司會議室",
    "present_directors": "李耿佑等 1 人出席",
    "registrationNumber": null,  // 設立後填寫
    "change_reason": "設立", // 設立/變更名稱/增資
    "date_of_authorization": "中華民國 113 年 04 月 22 日",
    "registrationType": "發起設立",
    "preexamination_Number": "113026425",
    "governmentFee": "1,000",
    "companyId": null, // 設立後填寫
    "year": 113,
    "shares_in_attendance": "100,000",
    "attendance_percentage": 100,
    "business_Scope": "1、J101010建築物清潔服務業。 2、JA03010洗衣業。 3、JZ99990未分類其他服務業。...",
    "Date_of_Adoption_of_the_Articles": "113年04月22日"
  },
  "table_data": {
    "board_minutes_table": [
      {
        "item": "1",
        "AGENDA_TITLE": "訂定公司章程案",
        "AGENDA_DESC": "依公司法第129條規定，擬定公司章程，如附章程。",
        "AGENDA_RESOLUTION": "經主席徵詢全體發起人無異議照案通過。"
      }
    ],
    "shareholders_minutes_table": [
      {
        "item": "1",
        "AGENDA_TITLE": "訂定公司章程案",
        "AGENDA_DESC": "依公司法第129條規定，擬定公司章程，如附章程。",
        "AGENDA_RESOLUTION": "經主席徵詢全體發起人無異議照案通過。"
      }
    ],
    "business_Scope_table": [
      {"code": "J101010", "description": "建築物清潔服務業"},
      {"code": "JA03010", "description": "洗衣業"},
      {"code": "JZ99990", "description": "未分類其他服務業"}
    ],
    "director_table": [
      {
        "title": "董事長",
        "name": "李耿佑",
        "id": "F128873285",
        "shares": 20000,
        "director_address_code":"110",
        "address": "新北市板橋區湳興里4鄰南雅西路二段7巷18之4號"
      },
      {
        "title": "監察人",
        "name": "李耿甫",
        "id": "A123456789",
        "shares": 0,
        "director_address_code":"110",
        "address": "台北市信義區忠孝東路一段1號"
      }
    ],
    "management_table": [
      // 經理人名單
    ],
    "leagalperson_table": [
      {
        "director_number": "3",
        "legal_entity_name": "建白企業有限公司",
        "legal_entity_id": "42783238",
        "legal_entity_address": "臺北市中山區長安東路2段81號11樓之4"
      }
    ],
    "shareholders_list": [
      {
        "item": "1",
        "shareholderName": "李耿佑",
        "shareholderId": "F128873285",
        "shareholderAddress": "新北市板橋區湳興里4鄰南雅西路二段7巷18之4號",
        "shareholder_shares": 20000,
        "shareholder_amount": 200000
      },
      {
        "item": "2",
        "shareholderName": "高淑敏",
        "shareholderId": "H223908974",
        "shareholderAddress": "台北市信義區永吉路225巷9號2樓",
        "shareholder_shares": 14500,
        "shareholder_amount": 145000
      },
      {
        "item": "3",
        "shareholderName": "建白企業有限公司",
        "shareholderId": "42783238",
        "shareholderAddress": "臺北市中山區長安東路2段81號11樓之4",
        "shareholder_shares": 65500,
        "shareholder_amount": 655000
      }
    ]
  }
}


# -----------------------------------------------------
# 核心函式 (C): 新增議程到表格 (統一複製，刪除樣板 - 修正重複問題)
# -----------------------------------------------------
def add_agendas_to_table(doc, header_tag, json_agendas):
    
    agenda_table = find_table_by_header(doc, header_tag)
    
    if agenda_table is None:
        print(f"❌ 錯誤：文件中找不到標題為 '{header_tag}' 的表格。")
        return

    # 樣板 Row 區塊的起始索引和數量
    # 假設標題是索引 0，樣板從索引 1, 2, 3 開始
    TEMPLATE_START_INDEX = 1
    ROW_COUNT_PER_AGENDA = 3 

    if len(agenda_table.rows) < TEMPLATE_START_INDEX + ROW_COUNT_PER_AGENDA:
        print("❌ 錯誤：議程表格中沒有足夠的行作為樣板 (需要至少 4 行)。")
        return
    
    print(f"✅ 找到議程表格。將複製樣板並填充 {len(json_agendas)} 個議程。")

    # 1. 遍歷 JSON 數據並複製/填充 (統一使用複製邏輯)
    for i, agenda_item in enumerate(json_agendas):
        
        agenda_num = i + 1 # 從 1, 2, 3... 開始編號
        
        # *** 核心步驟：複製樣板區塊 (每次都複製原始模板索引 1, 2, 3) ***
        # 不論 i=0 或 i>0，都複製，確保拿到的是「乾淨」的樣式結構
        target_rows = _copy_row_block(agenda_table, TEMPLATE_START_INDEX, ROW_COUNT_PER_AGENDA)
        
        # 替換映射
        replacements = {
            "{{item}}": str(agenda_num),
            "{{AGENDA_TITLE}}": agenda_item.get("案由：", ""),
            "{{AGENDA_DESC}}": agenda_item.get("說明：", ""),
            "{{AGENDA_RESOLUTION}}": agenda_item.get("決議：", "")
        }
        
        # 填充內容到 3 個 Row 區塊
        
        # Row 1 (索引 0): 案由, 編號
        _replace_in_paragraph(target_rows[0].cells[0].paragraphs[0], {"{{item}}": replacements["{{item}}"]})
        _replace_in_paragraph(target_rows[0].cells[2].paragraphs[0], {"{{AGENDA_TITLE}}": replacements["{{AGENDA_TITLE}}"]})
        
        # Row 2 (索引 1): 說明
        _replace_in_paragraph(target_rows[1].cells[2].paragraphs[0], {"{{AGENDA_DESC}}": replacements["{{AGENDA_DESC}}"]})

        # Row 3 (索引 2): 決議
        _replace_in_paragraph(target_rows[2].cells[2].paragraphs[0], {"{{AGENDA_RESOLUTION}}": replacements["{{AGENDA_RESOLUTION}}"]})

    
    # 2. 迴圈結束後，移除原始的樣板 Row 區塊 (索引 1, 2, 3)
    # 這是確保重複內容消失的關鍵步驟。
    for j in reversed(range(ROW_COUNT_PER_AGENDA)):
        row_index_to_remove = TEMPLATE_START_INDEX + j
        row_element = agenda_table.rows[row_index_to_remove]._element
        row_element.getparent().remove(row_element)
    
    print(f"✨ 議程 '{header_tag}' 填充完成，並已移除樣板區塊。")

    # 5. 如果 JSON 數據量 < 原始模板 Row 數量，則刪除多餘的樣板 Row
    # (此步驟可選，但為了保持程式碼簡潔性，暫時省略，通常我們會保留一個樣板)

# -----------------------------------------------------
# 終極穩健替換函數 (您提供的 _replace_in_paragraph 和 replace_all_placeholders)
# -----------------------------------------------------

def _replace_in_paragraph(para, mapping):
    # ... (您的 _get_font_settings 和 _set_font_settings 輔助函式放在這裡) ...
    # 內部輔助函數：從 Run XML 中獲取字體設定
    def _get_font_settings(run):
        settings = {}
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                settings['ascii'] = rFonts.get(qn('w:ascii'))
                settings['eastAsia'] = rFonts.get(qn('w:eastAsia'))
        return settings

    # 內部輔助函數：將字體設定套用到新的 Run XML
    def _set_font_settings(run, settings):
        if settings:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            
            if settings.get('ascii'):
                rFonts.set(qn('w:ascii'), settings['ascii'])
                rFonts.set(qn('w:hAnsi'), settings['ascii']) 
            
            if settings.get('eastAsia'):
                rFonts.set(qn('w:eastAsia'), settings['eastAsia'])

    full_text = para.text
    
    for placeholder, value in mapping.items():
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, str(value))
            
            style = None
            is_bold, is_italic, is_underline = None, None, None
            font_name, font_size = None, None
            cjk_font_settings = {}

            if para.runs:
                source_run = para.runs[0] 
                style = source_run.style 
                is_bold = source_run.bold
                is_italic = source_run.italic
                is_underline = source_run.underline
                
                if source_run.font.name: font_name = source_run.font.name
                if source_run.font.size: font_size = source_run.font.size
                    
                cjk_font_settings = _get_font_settings(source_run)

            while para.runs:
                run_element = para.runs[0]._element
                run_element.getparent().remove(run_element)
            
            new_run = para.add_run(new_text)
            
            if style: new_run.style = style
            if is_bold is not None: new_run.bold = is_bold
            if is_italic is not None: new_run.italic = is_italic
            if is_underline is not None: new_run.underline = is_underline
            
            if font_name: new_run.font.name = font_name
            if font_size: new_run.font.size = font_size

            _set_font_settings(new_run, cjk_font_settings)
            full_text = new_text
            
def replace_all_placeholders(doc, mapping):
    """遞歸替換文件中的佔位符 (調用穩健替換函數)"""
    # 1. 處理所有段落 (頂層)
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)

    # 2. 處理所有表格內的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, mapping)


# -----------------------------------------------------
# 主程式
# -----------------------------------------------------
def generate_document():
    
    # 1. 取得變更類型輸入
    user_input = input("請輸入變更類型 (多個類型請用逗號分隔): ")
    change_types = [t.strip() for t in user_input.split(",")]
    print(f"輸入的變更類型: {change_types}")

    # 2. 取得 Required_Documents 及 other_Documents
    required_docs = set()
    other_docs = set()
    
    reg_list = registration_doc.data.get("Company_Registration_Complete_List", [])
    
    for c_type in change_types:
        found = False
        for item in reg_list:
            if item["Registration_Type"] == c_type:
                found = True
                if "Required_Documents" in item:
                    required_docs.update(item["Required_Documents"])
                if "other_Documents" in item:
                    other_docs.update(item["other_Documents"])
                break
        if not found:
            print(f"⚠️ 警告: 找不到變更類型 '{c_type}' 的資料")

    # 3. Print 出來
    print("\n--- 應備文件 (Required Documents) ---")
    for doc in required_docs:
        print(f"- {doc}")
        
    print("\n--- 其他文件 (Other Documents) ---")
    for doc in other_docs:
        print(f"- {doc}")

    # 4. 動態產生文件
    print("\n--- 開始產生文件 ---")
    
    # 確保輸出目錄存在
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    reg = data["registration_data"]
    agenda_data = data.get("agend", [])
    AGENDA_HEADER_TAG = "{{board_minutes_table}}" 

    # 頂層文字替換資料
    replacements = {
        "{{companyName}}": reg.get("companyName", ""),
        "{{registrationNumber}}": reg.get("registrationNumber", ""),
        "{{zipcode}}": reg.get("zipcode", ""), 
        "{{address}}": reg.get("address", ""),
        "{{chairperson}}": reg.get("chairperson", ""),
    }

    for doc_name in required_docs:
        # 嘗試尋找對應的模板文件
        # 優先順序: .docx -> .doc
        template_path = None
        possible_filenames = [f"{doc_name}.docx", f"{doc_name}.doc"]
        
        for fname in possible_filenames:
            path = os.path.join(TEMPLATE_DIR, fname)
            if os.path.exists(path):
                template_path = path
                break
        
        if not template_path:
            print(f"⚠️ 跳過: 找不到 '{doc_name}' 的模板文件 (搜尋路徑: {TEMPLATE_DIR})")
            continue

        print(f"📄 處理模板: {template_path}")
        
        try:
            doc = Document(template_path)
            
            # 1. 頂層文字替換
            replace_all_placeholders(doc, replacements)
            
            # 2. 表格填充 (如果是董事會議事錄，且有議程資料)
            # 這裡簡單判斷，如果文件內有該 Tag 才執行
            # 或是根據 doc_name 判斷也可以，但 Tag 判斷較通用
            if find_table_by_header(doc, AGENDA_HEADER_TAG):
                 add_agendas_to_table(doc, AGENDA_HEADER_TAG, agenda_data)

            # 3. 存檔
            output_filename = f"{doc_name}_Output.docx"
            output_path = os.path.join(OUTPUT_DIR, output_filename)
            doc.save(output_path)
            print(f"   ✔ 已儲存: {output_path}")

        except Exception as e:
            print(f"   ❌ 處理失敗 '{doc_name}': {str(e)}")

    print("\n✨ 所有文件處理完成")


if __name__ == "__main__":
    generate_document()