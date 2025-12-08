"""
文件產生工具
- 根據 documents_required_list.json 中的 type 欄位自動判斷:
  - T: 使用範本 (templates 資料夾) + 變數/表格替換
  - H: 使用歷史資料 (history 資料夾) + 文字替換 + 儲存格更新
- 輸出檔案格式: 文件名稱_公司名稱.docx
"""
import json
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
import update_charter  # 儲存格內容替換模組


def replace_text_in_paragraph(paragraph, replacements):
    """替換段落中的文字"""
    for replacement in replacements:
        old_text = replacement.get('old_text', '')
        new_text = replacement.get('new_text', '')
        if old_text and old_text in paragraph.text:
            # 需要處理 runs 以保留格式
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)


def replace_text_in_table(table, replacements):
    """替換表格中的文字"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)


def replace_text_in_document(doc_path, replacements, output_path):
    """使用 python-docx 替換文件中的文字"""
    doc = Document(doc_path)
    
    # 替換段落中的文字
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # 替換表格中的文字
    for table in doc.tables:
        replace_text_in_table(table, replacements)
    
    # 替換頁首頁尾
    for section in doc.sections:
        # 頁首
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            replace_text_in_table(table, replacements)
        # 頁尾
        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            replace_text_in_table(table, replacements)
    
    doc.save(output_path)


def get_required_documents():
    """共用函數: 詢問變更種類並取得必要文件清單 (新格式: 包含 type 欄位)"""
    json_path = r"c:\Users\joe70\PythonProject\documentAI\documents_required_list.json"
    
    # 讀取 JSON 檔案
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"找不到檔案: {json_path}")
        return None
    except json.JSONDecodeError:
        print(f"檔案格式錯誤: {json_path}")
        return None

    # 詢問變更種類
    print("\n請輸入本次變更種類 (多種變更請用 ',' 分隔)")
    user_input = input("例如 '遷址(同縣市), 變更負責人': ").strip()

    # 處理輸入：支援全形逗號轉半形，並分割
    user_input = user_input.replace('，', ',')
    input_types = [t.strip() for t in user_input.split(',') if t.strip()]

    if not input_types:
        print("未輸入任何內容")
        return None

    # 建立 mapping (排除 default)
    registration_map = {}
    default_docs = []
    for item in data.get("documents_required_list", []):
        reg_type = item.get("Registration_Type")
        if reg_type == "default":
            default_docs = item.get("Required_Documents", [])
        else:
            registration_map[reg_type] = item.get("Required_Documents", [])

    # 依照取得的種類，取回需要文件(Required_Documents)並移除重複
    # 新格式: 每個文件是 {"name": "xxx", "type": "T/H"}
    all_required_docs = []
    seen_docs = set()
    found_types = []
    not_found_types = []

    for type_name in input_types:
        if type_name in registration_map:
            found_types.append(type_name)
            docs = registration_map[type_name]
            for doc in docs:
                # 支援新格式 (物件) 和舊格式 (字串)
                if isinstance(doc, dict):
                    doc_name = doc.get('name', '')
                    doc_type = doc.get('type', 'T')
                else:
                    doc_name = doc
                    doc_type = 'T'
                
                if doc_name and doc_name not in seen_docs:
                    seen_docs.add(doc_name)
                    all_required_docs.append({'name': doc_name, 'type': doc_type})
        else:
            not_found_types.append(type_name)

    # 如果完全沒有找到任何種類，使用 default
    if not found_types:
        print(f"\n找不到您輸入的所有種類: {', '.join(not_found_types)}")
        # 處理 default 文件
        for doc in default_docs:
            if isinstance(doc, dict):
                doc_name = doc.get('name', '')
                doc_type = doc.get('type', 'T')
            else:
                doc_name = doc
                doc_type = 'T'
            all_required_docs.append({'name': doc_name, 'type': doc_type})
        print("\n目前系統支援的種類如下:")
        for key in registration_map.keys():
            print(f"  - {key}")
    else:
        print(f"\n已識別的變更種類: {', '.join(found_types)}")
        if not_found_types:
            print(f"警告: 找不到以下種類: {', '.join(not_found_types)}")

    # 顯示必要文件清單
    print(f"\n合併後的必要文件清單 (已移除重複):")
    print("-" * 40)
    for doc in all_required_docs:
        type_label = "範本" if doc['type'] == 'T' else "歷史"
        print(f"[{doc['type']}] {doc['name']} ({type_label})")
    print("-" * 40)

    return {
        'all_required_docs': all_required_docs,
        'found_types': found_types,
        'not_found_types': not_found_types,
        'registration_map': registration_map
    }


def find_history_file(history_dir, doc_name):
    """
    在 history 資料夾中尋找符合的檔案
    檔案名稱格式可能是 "文件名稱_公司名稱.docx"
    比對時只看底線前的部分
    """
    docx_files = list(history_dir.glob("*.docx"))
    
    for file_path in docx_files:
        # 取得檔案名稱（不含副檔名）
        file_stem = file_path.stem
        # 取得底線前的部分作為文件名稱
        file_doc_name = file_stem.split('_')[0]
        
        if file_doc_name == doc_name:
            return file_path
    
    return None


def process_documents(data_content, all_required_docs, company_name):
    """處理所有文件 - 根據 type 欄位自動選擇處理方式"""
    base_dir = Path(r"c:\Users\joe70\PythonProject\documentAI")
    templates_dir = base_dir / "templates"
    history_dir = base_dir / "history"
    output_dir = base_dir / "output"
    
    # 確保輸出目錄存在
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 取得替換規則 (用於歷史資料模式)
    replacements = data_content.get('replacements', [])
    
    # 準備 context (用於範本模式)
    context = {}
    
    # 1. 將 basicInformation 展平放入根目錄
    if 'basicInformation' in data_content:
        context.update(data_content['basicInformation'])
    
    # 2. 處理表格資料 - 動態搜尋 tableData 內所有表格
    if 'tableData' in data_content:
        table_data = data_content['tableData']
        for table_name, table_content in table_data.items():
            if isinstance(table_content, dict):
                for key, value in table_content.items():
                    if isinstance(value, list):
                        context[key] = value
                    elif key not in ['document_title']:
                        context[key] = value
    
    # 3. 處理其他可能的表格資料（向下相容舊格式）
    for key, value in data_content.items():
        if key in ['basicInformation', 'tableData', 'replacements']:
            continue
        if isinstance(value, dict):
            for sub_key, sub_value in value.items():
                if isinstance(sub_value, list) and sub_key not in context:
                    context[sub_key] = sub_value
        elif isinstance(value, list) and key not in context:
            context[key] = value

    print(f"\n--- 開始處理 {len(all_required_docs)} 份文件 ---")
    
    template_success = 0
    template_skip = 0
    template_error = 0
    history_success = 0
    history_skip = 0
    history_error = 0

    for doc_info in all_required_docs:
        doc_name = doc_info['name']
        doc_type = doc_info['type']
        
        # 輸出檔名: 文件名稱_公司名稱.docx
        output_filename = f"{doc_name}_{company_name}.docx"
        output_path = output_dir / output_filename
        
        if doc_type == 'T':
            # 範本模式
            template_path = templates_dir / f"{doc_name}.docx"
            
            if not template_path.exists():
                print(f"[T-略過] 找不到範本: {doc_name}.docx")
                template_skip += 1
                continue
            
            try:
                doc = DocxTemplate(template_path)
                doc.render(context)
                doc.save(output_path)
                print(f"[T-成功] 已產生: {output_filename}")
                template_success += 1
            except Exception as e:
                print(f"[T-錯誤] 產生 {doc_name} 失敗: {e}")
                template_error += 1
                import traceback
                traceback.print_exc()
        
        elif doc_type == 'H':
            # 歷史資料模式
            history_file = find_history_file(history_dir, doc_name)
            
            if history_file is None:
                print(f"[H-略過] 在 history 資料夾中找不到: {doc_name}")
                history_skip += 1
                continue
            
            if not replacements:
                print(f"[H-警告] 沒有替換規則")
            
            try:
                # 步驟 1: 文字替換
                print(f"[H-處理] {doc_name} (來源: {history_file.name})")
                print(f"  步驟 1: 執行文字替換...")
                replace_text_in_document(history_file, replacements, output_path)
                
                # 步驟 2: 儲存格內容更新 (update_charter)
                print(f"  步驟 2: 檢查儲存格更新...")
                doc = Document(output_path)
                charter_results = update_charter.process_charter_updates(
                    doc, 
                    data_content, 
                    verbose=False  # 簡化輸出
                )
                
                # 如果有執行儲存格更新，則重新儲存
                if charter_results:
                    success_count = sum(1 for r in charter_results if r['success'])
                    if success_count > 0:
                        doc.save(output_path)
                        print(f"  - 已更新 {success_count} 個儲存格")
                    else:
                        print(f"  - 無符合的儲存格需要更新")
                else:
                    print(f"  - 沒有設定需要更新的儲存格")
                
                print(f"[H-成功] 已產生: {output_filename}")
                history_success += 1
                
            except Exception as e:
                print(f"[H-錯誤] 處理 {doc_name} 失敗: {e}")
                history_error += 1
                import traceback
                traceback.print_exc()
        
        else:
            print(f"[略過] 未知的類型 '{doc_type}': {doc_name}")

    # 顯示統計
    print(f"\n--- 處理完成 ---")
    print(f"範本模式 (T): 成功 {template_success} / 略過 {template_skip} / 錯誤 {template_error}")
    print(f"歷史模式 (H): 成功 {history_success} / 略過 {history_skip} / 錯誤 {history_error}")
    print(f"輸出目錄: {output_dir}")


def main():
    print("=" * 50)
    print("文件產生工具")
    print("=" * 50)
    
    # 讀取 data.json
    data_json_path = r"c:\Users\joe70\PythonProject\documentAI\data.json"
    try:
        with open(data_json_path, 'r', encoding='utf-8') as f:
            data_content = json.load(f)
    except Exception as e:
        print(f"讀取 data.json 失敗: {e}")
        return
    
    # 取得公司名稱
    company_name = data_content.get('basicInformation', {}).get('companyName', '未命名公司')
    print(f"\n公司名稱: {company_name}")
    
    # 取得必要文件清單 (包含 type 資訊)
    result = get_required_documents()
    if result is None:
        return
    
    all_required_docs = result['all_required_docs']
    found_types = result['found_types']
    
    # 如果沒有找到任何種類（使用了 default），就不產生文件
    if not found_types:
        print("\n由於無法判斷變更種類，不產生文件。")
        return
    
    # 處理所有文件
    process_documents(data_content, all_required_docs, company_name)


if __name__ == "__main__":
    main()
