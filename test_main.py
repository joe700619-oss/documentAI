import os
import json
from getbasicInformationfromMOEA import BasicInformationAPI
from get_required_documents import get_required_documents
import re
from dotenv import load_dotenv

from langchain_openai import ChatOpenAI
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# PDF / DOCX
from pypdf import PdfReader
import docx

# Unstructured
from unstructured.partition.auto import partition

load_dotenv()


# ==========================================================
# 文件類型定義與偵測
# ==========================================================
def detect_document_type(filename: str, text: str):
    name = filename.lower()
    head = text[:8000]

    if "變更" in name and "登記" in name:
        return "change_registration_form"

    if "設立" in name and "登記" in name:
        return "establishment_registration_form"

    if "董事" in name and "名單" in name:
        return "director_roster"

    if "股東" in name and "名單" in name:
        return "shareholder_roster"

    if "章程" in name:
        return "articles_of_association"

    if "會議" in name or "議事錄" in name:
        return "meeting_minutes"

    # content fallback
    if "變更登記" in head and head.count("董事") >= 3:
        return "change_registration_form"

    if "設立登記" in head and head.count("董事") >= 3:
        return "establishment_registration_form"

    if re.search(r"第.{1,3}條", head):
        return "articles_of_association"

    return "unknown"


def load_documents(folder):
    docs = []
    if not os.path.exists(folder):
        return docs

    for f in os.listdir(folder):
        path = os.path.join(folder, f)
        if not os.path.isfile(path):
            continue

        content = ""
        try:
            if path.endswith(".pdf") or path.endswith(".docx"):
                elements = partition(
                    filename=path,
                    strategy="hi_res",
                    max_characters=4000
                )
                content = "\n\n".join(str(el) for el in elements)
            else:
                with open(path, "r", encoding="utf-8", errors="ignore") as fh:
                    content = fh.read()
        except Exception:
            # fallback
            try:
                if path.endswith(".pdf"):
                    reader = PdfReader(path)
                    content = "\n".join(p.extract_text() for p in reader.pages if p.extract_text())
                elif path.endswith(".docx"):
                    doc = docx.Document(path)
                    content = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                continue

        if not content.strip():
            continue

        dtype = detect_document_type(f, content)
        docs.append(Document(
            page_content=content,
            metadata={"source": f, "doc_type": dtype}
        ))

    return docs

def parse_json_output(text):
    try:
        cleaned = text.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        if cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        return json.loads(cleaned.strip())
    except Exception as e:
        print(f"JSON Parse Error: {e}")
        return {}


# ==========================================================
# LLM 處理邏輯
# ==========================================================

def create_complete_old_state(llm, full_text: str, moea_data: dict, target_schema: dict):
    """
    Step 2 & 3: 產生完整的【舊資料狀態】(Old State)。
    邏輯：
    - 以 MOEA 資料為權威基礎。
    - 使用文件內容 (變更登記表/設立登記表) 補足 MOEA 缺少的資料 (例如董監事名單、營業項目細項)。
    - ⚠️ 注意：這一步驟【不應用】使用者的變更需求。我們要建立的是「變更前」的完整狀態。
      如果文件本身是「變更後」的登記表（歷史資料），則視為目前的「現狀」。
    """
    moea_info_str = json.dumps(moea_data, ensure_ascii=False, indent=2)
    schema_str = json.dumps(target_schema, ensure_ascii=False, indent=2)

    prompt = ChatPromptTemplate.from_template("""
你是一個資料整合專家。

任務：
請建立該公司目前的【完整現狀資料】(Old State JSON)。

來源資訊：
1. 【MOEA 權威資料】：官方基礎資料 (最準確)。
2. 【文件內容】：歷史文件，用於補足 MOEA 缺少的細節 (如董監事名單、詳細營業項目)。

整合規則：
1. **補足闕漏**：若 MOEA 資料有缺 (例如缺少 directors 列表)，請從【文件內容】中提取並補上。
2. **維持現狀**：請不要「更動」資料，除非 MOEA 是空的。我們要還原該公司在「本次變更前」的狀態。
3. **格式統一**：請依照【目標 JSON 格式】輸出。
4. **registration_type**：請留空或填寫 "current_status"，因為這是現狀資料。

【MOEA 權威資料】
{moea_info}

【文件內容 (歷史/現狀)】
{content}

【目標 JSON 格式】
{target_schema}

請直接回傳 JSON (不要包含 Markdown 標記)：
""")

    chain = prompt | llm | StrOutputParser()
    result = chain.invoke({
        "content": full_text,
        "moea_info": moea_info_str,
        "target_schema": schema_str
    })

    return parse_json_output(result)


def apply_changes_and_generate_new_state(llm, old_state: dict, user_request: str, allowed_types: list):
    """
    Step 4: 根據使用者需求，修改 Old State -> 產生 New State。
    同時判斷 registration_type。
    """
    old_state_str = json.dumps(old_state, ensure_ascii=False, indent=2)
    allowed_str = json.dumps(allowed_types, ensure_ascii=False, indent=2)

    prompt = ChatPromptTemplate.from_template("""
你是一個公司資料變更引擎。

任務：
請根據【使用者需求】，修改【舊資料狀態】，產生【新資料狀態】，並判斷【變更類型】。

輸入資訊：
1. 【舊資料狀態 (Old State)】：
{old_state}

2. 【使用者需求 (User Request)】：
{user_request}

3. 【合法變更類型列表】：
{allowed_types}

執行步驟：
1. **應用變更**：
   - 根據使用者需求修改舊資料中的對應欄位 (例如：姓名、地址、資本額)。
   - 若需求為「更換負責人」，請記得更新 basicInformation 下的 chairmanName 以及 tableData 中的 directors 列表。
   - 更新 `applicationReason` 欄位為使用者需求的摘要 (例如 "更換負責人為王大明")。

2. **判斷變更類型**：
   - 根據你所做的修改，從列表中選擇對應的 `registration_type`。
   - 若為遷址，請判斷縣市是否變更。
   - 將選定的類型填入 JSON 的 `registration_type` 欄位 (多選時用逗號分隔)。

3. **輸出結果**：
   - 回傳完整的【新資料狀態 JSON】。

請直接回傳 JSON (不要包含 Markdown 標記)：
""")

    chain = prompt | llm | StrOutputParser()
    result = chain.invoke({
        "old_state": old_state_str,
        "user_request": user_request,
        "allowed_types": allowed_str
    })

    return parse_json_output(result)


def compare_old_and_new(llm, old_state: dict, new_state: dict):
    """
    Step 5: 比對 Old 與 New，產生 changes_summary。
    使用 LLM 進行語意比對，避免格式差異造成的誤判。
    """
    old_str = json.dumps(old_state, ensure_ascii=False, indent=2)
    new_str = json.dumps(new_state, ensure_ascii=False, indent=2)

    prompt = ChatPromptTemplate.from_template("""
你是一個資料差異比對專家。

任務：
請比對【舊資料】與【新資料】，列出所有實質變更。

輸入資訊：
1. 【舊資料 (Old)】
{old_state}

2. 【新資料 (New)】
{new_state}

比對規則：
1. 忽略格式差異 (改行、空白)。
2. 忽略 "registration_type", "applicationReason", "date_of_adoption" 等本次申請必然變動的行政欄位，除非使用者明確要求變更。
3. **重點檢查**：
   - 公司名稱, 資本額, 地址, 負責人。
   - 董監事名單 (新增/移除/職位變更)。
   - 營業項目 (新增/移除)。

輸出格式：
請回傳一個 JSON Array (List), 每個項目如下：
{{
  "field": "欄位識別碼 (例如 companyName, directors)",
  "label": "中文欄位名稱",
  "old": "舊值 (摘要)",
  "new": "新值 (摘要)",
  "description": "說明差異 (例如 '負責人由 A 變更為 B')"
}}

若無實質變更，回傳 []。

請直接回傳 JSON List (不要包含 Markdown 標記)：
""")

    chain = prompt | llm | StrOutputParser()
    result = chain.invoke({
        "old_state": old_str,
        "new_state": new_str
    })

    return parse_json_output(result)


# ==========================================================
# 主流程
# ==========================================================
class DebugWorkflow:

    def __init__(self):
        self.history_dir = "history"
        self.history_cases_dir = "history_cases"
        self.api_key = os.getenv("OPENAI_API_KEY")

        self.llm = ChatOpenAI(
            model="gpt-4o",
            temperature=0,
            api_key=self.api_key
        )
        self.moea_api = BasicInformationAPI()

        # 載入合法變更類型
        self.allowed_registration_types = []
        try:
            with open("documents_required_list.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                for item in data.get("documents_required_list", []):
                    rtype = item.get("Registration_Type")
                    if rtype and rtype != "default":
                        self.allowed_registration_types.append(rtype)
        except Exception as e:
            print(f"Warning: Failed to load documents_required_list.json: {e}")

        # 載入目標 JSON Schema
        self.target_schema = {}
        try:
            with open("company_registration_form.json", "r", encoding="utf-8") as f:
                self.target_schema = json.load(f)
        except Exception as e:
            print(f"Warning: Failed to load company_registration_form.json: {e}")

    def run(self):
        # 1. 取得 MOEA 資料
        print("\n=== STEP 1: Fetching MOEA Data ===")
        # 先掃描文件找統編，若無則用預設
        all_docs = load_documents(self.history_dir) + load_documents(self.history_cases_dir)
        target_types = ["change_registration_form", "establishment_registration_form"]
        filtered_docs = [d for d in all_docs if d.metadata["doc_type"] in target_types]

        if not filtered_docs:
            print("未找到「變更登記表」或「設立登記表」，流程結束。")
            return

        full_text = "\n\n".join(d.page_content for d in filtered_docs)
        print(f"Read {len(filtered_docs)} doc(s).")
        
        tax_id_match = re.search(r"統一編號[：:\s]*(\d{8})", full_text)
        if not tax_id_match:
            tax_id_match = re.search(r"\b(\d{8})\b", full_text)
        tax_id = tax_id_match.group(1) if tax_id_match else "60299784"
        print(f"Target Tax ID: {tax_id}")

        moea_facts = self.moea_api.get_company_facts(tax_id)
        print("MOEA Data Fetched.")

        # 2. 歷史資料補足 -> 產生【完整舊資料】
        print("\n=== STEP 2: Creating Complete Old State (MOEA + Doc) ===")
        print("Filling gaps (e.g. directors) using document content...")
        
        old_state_json = create_complete_old_state(
            self.llm,
            full_text,
            moea_facts,
            self.target_schema
        )
        # Debug save
        with open("debug_old_state.json", "w", encoding="utf-8") as f:
            json.dump(old_state_json, f, ensure_ascii=False, indent=2)

        # 3. 使用者輸入
        print("\n=== STEP 3: User Request & Change Classification ===")
        user_input_request = input("\n請輸入您的變更需求 (例如：更換負責人為王大明、遷址...): ").strip()

        # 4. 產生【新資料】 (Old + Request)
        print("\n=== STEP 4: Generating New State (Applying Changes) ===")
        new_state_json = apply_changes_and_generate_new_state(
            self.llm,
            old_state_json,
            user_input_request,
            self.allowed_registration_types
        )

        # 5. 比對差異
        print("\n=== STEP 5: Generating Changes Summary ===")
        changes = compare_old_and_new(
            self.llm,
            old_state_json,
            new_state_json
        )
        
        print(f"Detected {len(changes)} changes.")
        for c in changes:
            print(f"- {c['label']} ({c['field']}): {c.get('description', '')}")

        # 6. Get Required Documents based on registration_type
        print("\n=== STEP 6: Getting Required Documents ===")
        registration_types = new_state_json.get("registration_type", "")
        # 處理可能是字串或列表的情況
        if isinstance(registration_types, str):
            # 如果是字串,根據逗號分割
            allowed_types = [t.strip() for t in registration_types.split(",") if t.strip()]
        elif isinstance(registration_types, list):
            allowed_types = registration_types
        else:
            allowed_types = []
        
        print(f"Registration Type(s): {allowed_types}")
        
        if allowed_types:
            required_docs = get_required_documents(allowed_types)
            print(f"Found {len(required_docs)} required documents:")
            for doc in required_docs:
                print(f"  - {doc['name']} (type: {doc['type']})")
            new_state_json["required_documents"] = required_docs
        else:
            print("No registration type found, skipping document retrieval.")
            new_state_json["required_documents"] = []

        # 7. Final Output
        print("\n=== Final JSON Output ===")
        new_state_json["changes_summary"] = changes
        
        output_file = "final_llm_output.json"
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(new_state_json, f, ensure_ascii=False, indent=2)
            
        print(json.dumps(new_state_json, ensure_ascii=False, indent=2))
        print(f"\n(Result saved to {output_file})")

if __name__ == "__main__":
    DebugWorkflow().run()
