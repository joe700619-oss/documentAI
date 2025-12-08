"""
update_charter.py
章程更新模組 - 提供儲存格內容替換功能

功能：
- 根據 data.json 中的 tableData 設定，找到文件中符合 target_start_text 的儲存格
- 將儲存格內容替換為 data.json 中定義的新內容

使用方式：
    from update_charter import process_charter_updates
    process_charter_updates(document, data_content)
"""
from docx import Document
from pathlib import Path
from docx.oxml.ns import qn


def format_replacement_content(table_config):
    """
    根據 table_config 的結構，格式化替換內容
    
    支援的資料類型:
    - business_items (營業項目): id, code, name
    - business_items (章程修訂日期): id, edition_date
    - shareholders (股東名冊): name, capital
    """
    target_text = table_config.get("target_start_text", "")
    
    # 判斷資料類型並格式化
    if "business_items" in table_config:
        items = table_config["business_items"]
        # 檢查是營業項目還是章程修訂日期
        if items and "code" in items[0]:
            # 營業項目格式: id. code name
            lines = [target_text]
            for item in items:
                lines.append(f"{item['id']}. {item['code']} {item['name']}")
            return "\n".join(lines)
        elif items and "edition_date" in items[0]:
            # 章程修訂日期格式
            lines = [target_text]
            for item in items:
                lines.append(f"{item['id']}. {item['edition_date']}")
            return "\n".join(lines)
    
    if "shareholders" in table_config:
        items = table_config["shareholders"]
        lines = [target_text]
        for item in items:
            lines.append(f"{item['name']}：{item['capital']}")
        return "\n".join(lines)
    
    return target_text


def find_and_replace_by_config(document, table_config, verbose=True):
    """
    根據 table_config 中的 target_start_text 找到儲存格並替換內容
    
    Args:
        document: python-docx Document 物件
        table_config: 包含 target_start_text 和替換資料的設定
        verbose: 是否輸出詳細訊息
    
    Returns:
        tuple: (success: bool, message: str)
    """
    target_start_text = table_config.get("target_start_text", "")
    if not target_start_text:
        return False, "No target_start_text defined"
    
    # 格式化替換內容
    new_content = format_replacement_content(table_config)
    
    # 遍歷所有表格
    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 移除冒號差異，做更寬鬆的比對
                target_clean = target_start_text.rstrip("：:")
                if cell_text.startswith(target_clean):
                    if verbose:
                        print(f"    [找到] 表格 {table_idx+1}, 列 {row_idx+1}, 格 {cell_idx+1}")
                        print(f"      原始內容: {cell_text[:50]}...")
                    
                    # 寫入新內容並設定字體
                    cell.text = ""  # 清空內容
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(new_content)
                    
                    # 設定字體為標楷體
                    run.font.name = '標楷體'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    
                    if verbose:
                        print(f"      已替換! (字體設定為標楷體)")
                    return True, "替換成功"
    
    return False, f"找不到以 '{target_start_text}' 開頭的儲存格"


def get_target_tables(data_content):
    """
    從 data_content 中取得所有含有 target_start_text 的表格設定
    
    Args:
        data_content: data.json 的內容
    
    Returns:
        list: 包含 name 和 config 的表格清單
    """
    table_data = data_content.get("tableData", {})
    target_tables = []
    
    for table_name, table_config in table_data.items():
        if "target_start_text" in table_config:
            target_tables.append({
                "name": table_name,
                "config": table_config
            })
    
    return target_tables


def process_charter_updates(document, data_content, verbose=True):
    """
    處理所有章程相關的儲存格更新
    
    Args:
        document: python-docx Document 物件
        data_content: data.json 的內容
        verbose: 是否輸出詳細訊息
    
    Returns:
        list: 包含每個替換結果的清單
    """
    target_tables = get_target_tables(data_content)
    
    if not target_tables:
        if verbose:
            print("    沒有找到需要替換的表格設定 (無 target_start_text)")
        return []
    
    if verbose:
        print(f"    找到 {len(target_tables)} 個表格設定:")
        for t in target_tables:
            print(f"      - {t['name']}: '{t['config']['target_start_text']}'")
    
    results = []
    for t in target_tables:
        if verbose:
            print(f"\n    處理 [{t['name']}]:")
        success, message = find_and_replace_by_config(document, t['config'], verbose)
        results.append({
            "name": t['name'],
            "success": success,
            "message": message
        })
    
    return results


def update_document_from_file(input_path, output_path, data_content, verbose=True):
    """
    讀取文件、執行更新、儲存文件
    
    Args:
        input_path: 輸入文件路徑
        output_path: 輸出文件路徑
        data_content: data.json 的內容
        verbose: 是否輸出詳細訊息
    
    Returns:
        tuple: (success: bool, results: list)
    """
    try:
        # 載入文件
        document = Document(input_path)
        
        # 執行替換
        results = process_charter_updates(document, data_content, verbose)
        
        # 確保輸出目錄存在
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 儲存文件
        document.save(output_path)
        
        return True, results
    
    except Exception as e:
        return False, [{"name": "Error", "success": False, "message": str(e)}]


# 測試用 main 函數
if __name__ == "__main__":
    import json
    
    base_dir = Path(r"c:\Users\joe70\PythonProject\documentAI")
    template_path = base_dir / "templates" / "公司章程.docx"
    output_path = base_dir / "output" / "公司章程_測試.docx"
    data_path = base_dir / "data.json"
    
    print("=" * 60)
    print("update_charter.py 測試")
    print("=" * 60)
    
    # 讀取 data.json
    with open(data_path, 'r', encoding='utf-8') as f:
        data_content = json.load(f)
    
    # 執行更新
    success, results = update_document_from_file(
        template_path, 
        output_path, 
        data_content, 
        verbose=True
    )
    
    print("\n" + "=" * 60)
    print("結果:")
    print("=" * 60)
    for r in results:
        status = "OK" if r['success'] else "SKIP"
        print(f"  [{status}] {r['name']}: {r['message']}")
    
    print(f"\n輸出: {output_path}")
